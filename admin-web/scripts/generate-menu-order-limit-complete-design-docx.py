# -*- coding: utf-8 -*-
"""Generate full 数量与频次限制完整设计方案 (MD + DOCX) aligned to template structure."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs/superpowers/specs/menu-order-limit"
OUT_MD = OUT_DIR / "2026-08-20-menu-order-limit-complete-design.md"
OUT_DOCX = OUT_DIR / "2026-08-20-menu-order-limit-complete-design.docx"
DL_DOCX = Path(r"c:\Users\27273\Downloads\数量与频次限制完整设计方案-项目落地版.docx")


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


class Writer:
    def __init__(self) -> None:
        self.md: list[str] = []
        self.doc = Document()
        set_doc_font(self.doc)
        for s in self.doc.sections:
            s.top_margin = Inches(1)
            s.bottom_margin = Inches(1)
            s.left_margin = Inches(1)
            s.right_margin = Inches(1)

    # --- markdown ---
    def m(self, line: str = "") -> None:
        self.md.append(line)

    def m_h1(self, t: str) -> None:
        self.m(f"\n# {t}\n")

    def m_h2(self, t: str) -> None:
        self.m(f"\n## {t}\n")

    def m_h3(self, t: str) -> None:
        self.m(f"\n### {t}\n")

    def m_p(self, t: str) -> None:
        self.m(t)

    def m_bullets(self, items: list[str]) -> None:
        for i in items:
            self.m(f"- {i}")

    def m_table(self, headers: list[str], rows: list[list[str]]) -> None:
        self.m("| " + " | ".join(headers) + " |")
        self.m("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows:
            self.m("| " + " | ".join(row) + " |")
        self.m("")

    def m_code(self, text: str) -> None:
        self.m("```text")
        self.m(text)
        self.m("```\n")

    def m_status(self, label: str) -> None:
        self.m(f"> **落地状态：{label}**\n")

    # --- docx ---
    def d_title(self, t: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.bold = True
        r.font.size = Pt(22)

    def d_h1(self, t: str) -> None:
        self.doc.add_heading(t, level=1)

    def d_h2(self, t: str) -> None:
        self.doc.add_heading(t, level=2)

    def d_h3(self, t: str) -> None:
        self.doc.add_heading(t, level=3)

    def d_p(self, t: str, bold: bool = False) -> None:
        p = self.doc.add_paragraph()
        r = p.add_run(t)
        r.bold = bold

    def d_status(self, label: str) -> None:
        p = self.doc.add_paragraph()
        r = p.add_run(f"落地状态：{label}")
        r.bold = True
        r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    def d_bullets(self, items: list[str]) -> None:
        for i in items:
            self.doc.add_paragraph(i, style="List Bullet")

    def d_table(self, headers: list[str], rows: list[list[str]]) -> None:
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.shared import Inches as In

        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                t.rows[ri + 1].cells[ci].text = val
        self.doc.add_paragraph("")

    def d_code(self, text: str) -> None:
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Consolas"
        r.font.size = Pt(9)

    def both_h1(self, t: str) -> None:
        self.m_h1(t)
        self.d_h1(t)

    def both_h2(self, t: str) -> None:
        self.m_h2(t)
        self.d_h2(t)

    def both_h3(self, t: str) -> None:
        self.m_h3(t)
        self.d_h3(t)

    def both_p(self, t: str) -> None:
        self.m_p(t)
        self.d_p(t)

    def both_status(self, label: str) -> None:
        self.m_status(label)
        self.d_status(label)

    def both_bullets(self, items: list[str]) -> None:
        self.m_bullets(items)
        self.d_bullets(items)

    def both_table(self, headers: list[str], rows: list[list[str]]) -> None:
        self.m_table(headers, rows)
        self.d_table(headers, rows)

    def both_code(self, text: str) -> None:
        self.m_code(text)
        self.d_code(text)


def build(w: Writer) -> None:
    w.d_title("数量与频次限制完整设计方案")
    w.both_h1("菜单下单限制：数量与频次限制完整设计方案")
    meta = [
        "状态：已确认并已实现（后台配置原型；运行时判定为目标态）",
        "日期：2026-08-20",
        "模块：前厅管理中心 / 菜单下单限制 / 数量与频次限制",
        "版本：v1.1（与 PRD/SPEC MOL-01～MOL-31、专项 specs 对齐）",
        "方案：统一规则模型、统一计算语义；配置侧已落地，服务端最终裁决为目标态",
        "代码依据：order-limit-flow.js、order-limit.html、foh-menu-order-limits-ui.ts",
    ]
    for line in meta:
        w.both_p(line)

    # 1
    w.both_h1("1. 背景与目标")
    w.both_status("配置侧已落地；订单运行时判定为目标态")
    w.both_p(
        "前厅管理中心「菜单下单限制 > 数量与频次限制」已完成规则设计器、人数/轮次场景、"
        "按门店分类/菜品数量矩阵、生效范围（活动周期/营业时段/会员/门店）、超限授权配置、"
        "规则列表与六步发布确认。规则保存在浏览器 localStorage，尚未接入 POS/Kiosk/eMenu/SDI "
        "订单事务与服务端统一判定引擎。"
    )
    w.both_p(
        "本方案「频次」指每轮、多轮、与轮次无关三种统计周期，不包含「两次提交间隔秒数」等时间节流"
        "（属于「其他设置」Tab，不在本文范围）。"
    )
    w.both_h2("1.1 核心目标")
    w.both_bullets([
        "统一 12 种限购主体×统计周期×限购对象组合的配置、展示与计算公式解释。",
        "按人数限购 = 人均配置 L × 有效人数 N，不追踪具体食客。",
        "统一 storeConfigs[storeId] 数据模型：structureByLine、limits 数量矩阵。",
        "六步向导：规则类型→场景配置→限购数量→超限授权→生效范围→确认发布。",
        "商品选品并入限购数量；完整商品规则列表一行一条（门店×人数×轮次×产线×商品）。",
        "草稿 900ms 自动保存；sessionStorage 恢复副本仅写入/清理，不读取恢复。",
        "发布前全量校验 + 独立发布确认页；正式规则编辑派生草稿，发布后替换源 ID。",
        "数量三态：空=未配置、0=禁止、正整数=上限；不支持显式「不限制」。",
        "二次确认一律页面内自定义对话框（非 alert/confirm）。",
    ])
    w.both_h2("1.2 非目标")
    w.both_bullets([
        "首期不建立独立远程规则微服务；运行时引擎随订单服务部署（未落地）。",
        "不追踪每份商品属于哪位食客。",
        "首期不支持离线点单写入或运行时超限授权（未落地）。",
        "数量授权不绕过售罄、停售、支付等其他限制（运行时未接入）。",
        "不自动将旧 localStorage 发布为正式服务端规则；仅本地原型存储。",
        "隐藏的点菜模拟验证、其他两个 Tab 内部业务不在交付范围。",
    ])

    # 2
    w.both_h1("2. 术语")
    w.both_table(["术语", "定义", "配置/运行时"], [
        ["按订单限购", "subject=order；EffectiveLimit=L", "已配置 / 目标运行时"],
        ["按人数限购", "subject=party_size；EffectiveLimit=L×N", "已配置 / 目标运行时"],
        ["每轮", "period=per_round；各轮独立数量池", "已配置 / 目标运行时"],
        ["多轮", "period=multi_round；轮次区间不同 L", "已配置 / 目标运行时"],
        ["与轮次无关", "period=order_lifetime；全单累计", "已配置 / 目标运行时"],
        ["按分类限购", "targetType=category；分类内共享池", "已配置 / 目标运行时"],
        ["按菜品限购", "targetType=dish；每菜独立池", "已配置 / 目标运行时"],
        ["有效人数 N", "adult+child 或仅 adult（childCountPolicy）", "已配置 / 目标运行时"],
        ["Committed", "已提交且未释放占用量", "目标运行时"],
        ["CandidateCart", "应用本次变更后的候选数量", "目标运行时"],
        ["本次操作", "authorization scope=operation", "已配置 / 目标运行时"],
        ["当前轮授权", "scope=round", "已配置 / 目标运行时"],
        ["当前订单授权", "scope=order", "已配置 / 目标运行时"],
        ["参与门店", "participatingStoreIds；有 targetIds 的门店", "已落地"],
        ["生效门店", "deployStoreIds；发布快照门店", "已落地"],
    ])

    # 3
    w.both_h1("3. 总体架构")
    w.both_h2("3.1 当前落地架构（已实现）")
    w.both_code(
        "应用壳 /operations/queue-call/menu-order-limits\n"
        "  foh-menu-order-limits-ui.ts（三页签 + iframe 全屏）\n"
        "    order-limit.html（列表 + 筛选 + 字段设置）\n"
        "    order-limit-rule-editor.html（六步编辑器）\n"
        "    order-limit-publish-confirm.html（发布确认）\n"
        "      order-limit-flow.js\n"
        "        localStorage.restaurantRules\n"
        "        sessionStorage.restaurantRuleRecovery:{draftId}"
    )
    w.both_h2("3.2 目标态架构（模板原方案，未落地）")
    w.both_code(
        "后台设计器 → 规则管理 API → 门店规则存储 → 统一限购判定引擎\n"
        "POS/Kiosk/eMenu/SDI → 客户端预校验 → 判定引擎 → 加购/提交\n"
        "超限 → 服务员授权 → 凭证 → 重新判定 → 审计"
    )
    w.both_h2("3.3 架构原则")
    w.both_table(["原则", "配置侧现状", "目标态"], [
        ["统一 DTO/语义", "EditorDraft + 兼容顶层投影", "服务端/终端共用"],
        ["客户端预校验", "无", "及时提示剩余额度"],
        ["服务端最终裁决", "无", "写订单事务内校验"],
        ["规则版本", "本地覆盖写，无 version 字段", "不可变版本 + 回滚"],
        ["多终端共享池", "无", "同一 orderId 共享额度"],
        ["多规则 AND", "可配多条规则", "运行时 AND 判定"],
    ])

    # 4
    w.both_h1("4. 统一计算模型")
    w.both_p("定义：N=有效就餐人数；L=当前人数/轮次区间、分类或菜品配置数量。")
    w.both_code(
        "按订单：EffectiveLimit = L\n"
        "按人数：EffectiveLimit = L × N\n"
        "允许增加（目标态）：Committed + CandidateCart <= EffectiveLimit"
    )
    w.both_p(
        "配置侧在步骤 1 实时展示规则解释与公式；运行时 Committed/CandidateCart 由订单服务维护（未落地）。"
    )
    w.both_h2("4.1 数量三态（已确认）")
    w.both_table(["配置状态", "configured", "value", "UI 含义", "校验"], [
        ["留空/未配置", "false 或缺失", "null", "placeholder「未配置」", "不能离开步骤 3/发布"],
        ["禁止", "true", "0", "明确禁止", "已配置，可发布"],
        ["正整数", "true", ">0", "限购上限", "已配置，可发布"],
    ])
    w.both_p(
        "已移除：configured=true,value=null（旧「不限制」）→ 归一化为未配置。"
        "已移除：批量工具栏单独「设为禁止」按钮；输入 0 后「应用数量」表达禁止。"
    )

    # 5
    w.both_h1("5. 十二种场景")
    w.both_table(["场景", "实际限额", "统计范围", "配置"], [
        ["按订单+每轮+分类", "L", "当前订单、当前轮、分类总数", "已支持"],
        ["按订单+每轮+菜品", "L", "当前订单、当前轮、菜品总数", "已支持"],
        ["按订单+多轮+分类", "当前轮次区间 L", "当前订单、当前轮、分类总数", "已支持"],
        ["按订单+多轮+菜品", "当前轮次区间 L", "当前订单、当前轮、菜品总数", "已支持"],
        ["按订单+与轮次无关+分类", "L", "订单全部轮次、分类总数", "已支持"],
        ["按订单+与轮次无关+菜品", "L", "订单全部轮次、菜品总数", "已支持"],
        ["按人数+每轮+分类", "L×N", "当前订单、当前轮、分类总数", "已支持"],
        ["按人数+每轮+菜品", "L×N", "当前订单、当前轮、菜品总数", "已支持"],
        ["按人数+多轮+分类", "轮次区间 L×N", "当前订单、当前轮、分类总数", "已支持"],
        ["按人数+多轮+菜品", "轮次区间 L×N", "当前订单、当前轮、菜品总数", "已支持"],
        ["按人数+与轮次无关+分类", "L×N", "订单全部轮次、分类总数", "已支持"],
        ["按人数+与轮次无关+菜品", "L×N", "订单全部轮次、菜品总数", "已支持"],
    ])
    w.both_h2("5.1 分类规则示例")
    w.both_p("分类 A 含 a,b,c；按人数每轮 L=2，N=4 → 当前轮实际限额 2×4=8 份。")
    w.both_h2("5.2 菜品规则示例")
    w.both_p("四人按人数每轮：a 1/人→4；b 2/人→8；c 3/人→12。a 达上限只限 a。")
    w.both_h2("5.3 多轮规则示例")
    w.both_p("分类 A 按订单多轮：第1轮3、第2轮2、第3轮1、第4轮及以后0；各轮独立池。")

    # 6
    w.both_h1("6. 多规则叠加和重复校验")
    w.both_status("配置侧允许多规则；运行时 AND 与冲突检测为目标态")
    w.both_bullets([
        "不同维度可叠加：分类总量+单品、按订单+按人数、每轮+整单等（目标态 AND）。",
        "禁止重复：三核心维度完全相同 + 目标交集 + 门店/产线/时间/会员交集 + 同型无优先级（目标态发布校验）。",
        "当前原型：列表可存多条规则；发布时不做自动冲突检测；无服务端判定。",
    ])

    # 7 - six steps mapped to template 7.1-7.7
    w.both_h1("7. 后台配置流程")
    w.both_status("已落地（六步；原模板七步中「基础信息+商品」已合并）")
    w.both_h3("7.0 规则列表（入口页，order-limit.html）")
    w.both_table(["能力", "说明"], [
        ["新增规则", "创建 draft，进入步骤1"],
        ["编辑", "draft 继续；正式规则派生 sourceRuleId 草稿"],
        ["复制", "独立 draft，名称加「(副本)」"],
        ["查看", "ruleId + view=1 只读六步"],
        ["启停", "active↔inactive；draft 无"],
        ["删除", "自定义确认框"],
        ["筛选", "门店/状态/人数/轮次/时间 AND；重置"],
        ["字段设置", "分组显隐列；固定 name/status/actions"],
        ["滚动", "标题+表头固定，.section-body 内滚动"],
    ])
    w.both_h3("7.1 第一步：规则类型")
    w.both_table(["字段", "类型", "必填", "说明"], [
        ["name", "string", "是", "规则名称，maxlength=60"],
        ["description", "string", "否", "规则描述，maxlength=200"],
        ["subject", "order|party_size", "是", "限购主体"],
        ["period", "per_round|multi_round|order_lifetime", "是", "统计周期/频次"],
        ["targetType", "category|dish", "是", "限购对象；切换清空已选商品"],
        ["conditions.childCountPolicy", "inherit|include|exclude", "条件", "仅 party_size 展示"],
    ])
    w.both_p("实时展示规则类型解释与计算公式。校验：四维必选 + 名称非空。")

    w.both_h3("7.2 第二步：场景配置（原模板第三步）")
    w.both_table(["字段", "类型", "约束"], [
        ["partyRanges[]", "Range{min,max}", "从1连续、末段及以上、可自动补全"],
        ["roundRanges[]", "Range{min,max}", "仅 multi_round；同上连续规则"],
    ])
    w.both_p("区间变更（含自动补全）清空全部门店 limits；删除区间前确认。")

    w.both_h3("7.3 第三步：限购数量（合并原第二步商品范围 + 第四步数量）")
    w.both_bullets([
        "全局「添加商品」→ 门店 + 产线 + 分类/菜品；临时 productAddDialog 差异提交。",
        "storeConfigs[storeId].structureByLine / targetIds / limits。",
        "完整规则列表：筛选（门店/人数/轮次/产线/状态/搜索）、分页 10/20/50/100、批量填数/删商品。",
        "跨场景复制、跨产线复制、已选预览、已配数量预览。",
        "已移除：场景矩阵主区、产线 Tab 平铺、独立商品配置步骤、批量「设为禁止」/「设为不限制」。",
        "已移除模板四附加字段：每轮最大总份数、最大生效轮次、每人每轮上限、单次下单上限。",
    ])
    w.both_p("商品弹层 productAddDialog 字段：open, storeId, structureByLine, dirty, query, searchComposing。")
    w.both_p("规则行展开：门店 × partyRanges ×（multi_round ? roundRanges : 1）× 产线 × target → 一行。")
    w.both_p("规则行字段：勾选、配置门店、人数场景、轮次（多轮/每轮/与轮次无关文案）、产线、菜单、限购数量、移除。")
    w.both_table(["limitRuleList 字段", "默认", "行为"], [
        ["storeId", '""', "门店筛选；变更回第1页清勾选"],
        ["partyKey / roundKey", '""', "人数/轮次筛选"],
        ["lineId", '""', "产线筛选 kiosk|emenu|sdi"],
        ["status", '""', "configured|unconfigured"],
        ["query", '""', "菜品/分类名称搜索"],
        ["page / pageSize", "1 / 20", "10|20|50|100"],
        ["selectedRowIds", "[]", "跨页保留；筛选重置清空"],
    ])
    w.both_p("批量：勾选 selectedRowIds 后输入数量点「应用数量」；0=禁止；空输入校验不应用。"
             "批量删除按门店+产线+商品去重；有数量时一次汇总确认。")

    w.both_h3("7.4 第四步：超限授权（原模板第六步）")
    w.both_table(["字段", "类型", "默认", "说明"], [
        ["authorization.enabled", "boolean", "true", "false=硬性拒绝"],
        ["allowedScopes", "operation|round|order[]", "全开", "至少一项"],
        ["defaultScope", "scope", "round", "须属于 allowedScopes"],
        ["scopePermissions", "Record<scope,role>", "各范围角色", "值班经理/主管/店长/区域经理"],
        ["reasonRequired", "boolean", "true", "授权原因必填（配置项）"],
    ])

    w.both_h3("7.5 第五步：生效范围（原模板第五步）")
    w.both_table(["字段", "类型", "说明"], [
        ["effectiveFrom / effectiveTo", "date", "结束空=长期；结束≥开始"],
        ["activityCycle", "daily|weekly|monthly", "活动周期"],
        ["daysOfWeek", "weekday[]", "weekly 至少一天"],
        ["daysOfMonth", "1-31[]", "monthly 至少一日"],
        ["businessHourSlots", "{id,mode,from,to}[]", "all|lunch|dinner；full|custom"],
        ["businessHourSetupMode", "all_full|per_slot", "全时段/逐时段"],
        ["memberMode / memberLevelIds", "all|specified", "指定会员非空"],
        ["deployStoreIds", "string[]", "须为已添加商品门店，≥1"],
        ["deployExcludedStoreIds", "string[]", "用户主动取消生效的记忆"],
    ])
    w.both_p("营业时段边界：全天 00:00-23:59；午市 11:00-16:59；晚市 17:00-23:00。custom 须在父边界内。")
    w.both_table(["businessHourSlots[]", "类型", "说明"], [
        ["id", "all|lunch|dinner", "全天与午/晚互斥"],
        ["mode", "full|custom", "该时段全段或指定时间"],
        ["from / to", "HH:MM", "仅 custom；开始<结束；在父边界内"],
    ])
    w.both_p("activityCycle=daily 无额外日控件；weekly 用 daysOfWeek；monthly 用 daysOfMonth（当月无则跳过）。")

    w.both_h3("7.6 第六步：确认与发布（原模板第七步）")
    w.both_bullets([
        "汇总：规则公式、商品范围、人数/轮次矩阵完成度、生效条件、授权策略、生效门店。",
        "失败项「前往修正」跳转对应步骤。",
        "「保存并下发」→ order-limit-publish-confirm.html 二次 validateAll。",
    ])

    w.both_h3("7.7 生命周期")
    w.both_code("draft → active ↔ inactive\n编辑正式规则：sourceRuleId 草稿 → 发布替换源 ID\n复制：独立 draft，无 sourceRuleId")
    w.both_p("当前无 pending/archived 状态；启停仅正式规则。")

    # 8 full data model
    w.both_h1("8. 规则数据模型")
    w.both_status("已落地（localStorage JSON）")
    w.both_h2("8.1 StoredRule")
    w.both_table(["字段", "类型", "必填", "默认", "说明"], [
        ["id", "number|string", "是", "maxId+1", "规则 ID"],
        ["sourceRuleId", "id|null", "否", "null", "编辑正式规则来源"],
        ["status", "draft|active|inactive", "是", "draft", "无 pending/archived"],
        ["created", "YYYY-MM-DD", "是", "当日", "发布覆盖保留源 created"],
        ["publishedAt", "ISO", "正式", "发布时", "仅展示"],
        ["editorDraft", "EditorDraft", "是", "defaultDraft()", "运行/编辑快照"],
        ["authoringDraft", "EditorDraft", "正式", "发布保存", "编辑正式时优先读"],
        ["兼容顶层字段", "派生", "是", "buildCompatibilityRule", "列表/旧消费者"],
    ])

    w.both_h2("8.2 EditorDraft")
    w.both_table(["字段", "类型", "必填", "默认", "说明"], [
        ["currentStep / highestStep", "1-6", "是", "1", "步骤导航"],
        ["subject/period/targetType", "enum|null", "是", "null", "步骤1"],
        ["name/description", "string", "名称必填", '""', "步骤1"],
        ["partyRanges/roundRanges", "Range[]", "是", "[{1,null}]", "步骤2"],
        ["conditions", "Conditions", "是", "见8.4", "步骤5"],
        ["authorization", "Authorization", "是", "见8.5", "步骤4"],
        ["participatingStoreIds", "string[]", "是", "[]", "归一化=已添加门店"],
        ["storeConfigs", "Record<storeId,StoreConfig>", "是", "{}", "权威数据"],
        ["deployStoreIds", "string[]", "是", "[]", "生效门店"],
        ["deployExcludedStoreIds", "string[]", "是", "[]", "排除记忆"],
        ["productQuantityMergedVersion", "number", "迁移", "2", "步骤结构版本"],
        ["legacyCompatibilityFallback", "StoreConfig", "是", "空", "旧数据留存"],
    ])

    w.both_h2("8.3 StoreConfig 与 LimitCell")
    w.both_table(["字段", "类型", "说明"], [
        ["structureByLine", "{kiosk,emenu,sdi:[]}", "选品树"],
        ["productLines", "string[]", "有目标的产线"],
        ["targetIds", "string[]", "限购目标 ID"],
        ["limits", "Record<LimitKey,LimitCell>", "数量矩阵"],
        ["LimitKey", "party|round|line|targetId", "targetId 可含 |"],
        ["LimitCell.configured", "boolean", "是否显式配置"],
        ["LimitCell.value", "0|正整数|null", "null=未配置"],
    ])

    w.both_h2("8.4 Conditions 完整字段")
    w.both_table(["字段", "类型", "默认", "校验"], [
        ["effectiveFrom", "date", "当日", "≤ effectiveTo"],
        ["effectiveTo", "date|\"\"", "\"\"", "可空=长期"],
        ["activityCycle", "daily|weekly|monthly", "weekly", "周/月非空"],
        ["daysOfWeek", "weekday[]", "全周", "weekly"],
        ["daysOfMonth", "1-31[]", "[]", "monthly 去重升序"],
        ["businessHourSlots", "slot[]", "晚市 full", "≥1 时段"],
        ["businessHourSetupMode", "all_full|per_slot", "all_full", "非法值推导"],
        ["memberMode", "all|specified", "all", "specified 选等级"],
        ["memberLevelIds", "string[]", "[]", "specified 非空"],
        ["childCountPolicy", "inherit|include|exclude", "inherit", "步骤1 按人数"],
        ["businessHour* 旧字段", "兼容", "镜像", "由 slots 同步"],
    ])

    w.both_h2("8.5 Authorization 完整字段")
    w.both_table(["字段", "类型", "默认", "校验"], [
        ["enabled", "boolean", "true", "false=汇总硬性拒绝"],
        ["allowedScopes", "scope[]", "三种", "enabled 时 ≥1"],
        ["defaultScope", "scope|\"\"", "round", "∈ allowedScopes"],
        ["scopePermissions", "Record", "静态角色", "每启用范围非空"],
        ["reasonRequired", "boolean", "true", "配置项"],
    ])

    w.both_h2("8.6 界面临时状态（不持久化）")
    w.both_table(["对象", "字段", "说明"], [
        ["limitRuleList", "storeId,partyKey,roundKey,lineId,status,query,page,pageSize,selectedRowIds", "数量列表 UI"],
        ["productAddDialog", "open,storeId,structureByLine,dirty,query", "添加商品弹层"],
        ["列偏好", "version,visible[]", "order-limit:rule-list-columns:v1"],
    ])

    w.both_h2("8.7 规则列表列定义（已落地）")
    w.both_p("固定列：name, status, actions。默认可见：name, strategy, persons, productScope, "
             "effectiveStores, effectiveTime, authorization, status, actions。")
    w.both_p("可选列按六步分组：描述、主体、周期、对象、儿童口径、人数/轮次区间、参与门店、产线、"
             "目标数、完成度、授权范围/默认值/权限/原因、日期/周期/营业时段/会员。")
    w.both_p("列表筛选（AND）：门店、状态(draft/active/inactive)、人数场景、轮次、时间组合；无规则名称关键词筛选。")

    w.both_h2("8.8 六步渲染与校验映射")
    w.both_table(["步骤", "渲染函数", "离开校验"], [
        ["1 规则类型", "renderStepOne", "主体/周期/对象/名称"],
        ["2 场景配置", "renderStepThree", "区间连续+末段及以上"],
        ["3 限购数量", "renderStepFour", "≥1目标；全部 limits configured"],
        ["4 超限授权", "renderStepSix", "范围/默认/每范围权限"],
        ["5 生效范围", "renderStepFive", "门店/周期/会员/日期/时段"],
        ["6 确认发布", "renderStepSeven + 发布页", "validateAll + validateDeployStores"],
    ])

    w.both_h2("8.9 目标态 OrderLimitRule（模板原 TypeScript，未落地）")
    w.both_p("服务端正式版可增加 version、storeIds、menuIds、orderTypes、scenes 扁平结构、"
              "createdBy/updatedBy 等；当前原型以 EditorDraft 为准，发布时 buildPublishedDraft 裁剪门店。")

    # 9-14 target state with config cross-ref
    w.both_h1("9. 判定接口")
    w.both_status("目标态（模板）；当前无 HTTP 接口")
    w.both_h3("9.1 OrderLimitWriteRequest / OrderLimitContext")
    w.both_p("客户端只提交 orderId、operationId、orderVersion、menuId、mutation、authorizationTokens；"
             "人数/轮次/已下单/购物车由服务端读取。当前原型不实现。")
    w.both_h3("9.2 OrderLimitDecision")
    w.both_p("violations 含 ruleId、targetId、configuredLimit、effectiveLimit、各 quantity、exceededQuantity、"
             "allowedAuthorizationScopes、message。")
    w.both_h3("9.3 判定顺序")
    w.both_bullets([
        "加载门店正式规则→过滤门店/产线/菜单/订单类型→日期/时段/会员→有效人数→区间匹配→"
        "展开分类→统计 Committed+Candidate→计算 EffectiveLimit→校验授权→返回全部违反项。",
    ])

    w.both_h1("10. 轮次与数量占用")
    w.both_status("目标态")
    w.both_h3("10.1 轮次边界")
    w.both_bullets(["新单第1轮", "提交成功完成轮次", "失败不增轮", "重试保持 operationId"])
    w.both_h3("10.2 释放限额")
    w.both_table(["操作", "释放"], [
        ["未提交删购物车", "是"], ["提交前取消", "是"], ["送厨前撤销", "是"],
        ["已送厨退菜", "否"], ["提交失败", "不占用"],
    ])
    w.both_p("OrderItemLimitSnapshot / occupancies 按规则+目标逐条保存（目标态）。")

    w.both_h1("11. 人数变化")
    w.both_status("配置已落地 childCountPolicy；运行时重算为目标态")
    w.both_code("包含儿童：adultCount+childCount\n排除儿童：adultCount")
    w.both_p("人数区间在步骤2配置；运行时人数变更重算限额、不追溯删历史（目标态）。")

    w.both_h1("12. 服务员密码授权")
    w.both_h2("12.1 三种范围（配置已落地）")
    w.both_table(["范围", "配置字段", "目标态生效", "目标态失效"], [
        ["本次操作", "operation", "当前请求", "成功后消费"],
        ["当前轮", "round", "当前轮内免重复验密", "换轮/关单"],
        ["当前订单", "order", "订单内免重复验密", "关单/规则版本变"],
    ])
    w.both_h2("12.2 授权流程（目标态）")
    w.both_bullets(["超限→拒绝写入→选范围→密码→签凭证→重试→事务写入→审计"])
    w.both_h2("12.3 OrderLimitAuthorization 凭证（目标态）")
    w.both_p("authorizationId、ruleRefs、scope、operationId、requestDigest、roundNo、expiresAt、status…")

    w.both_h1("13. 客户端交互")
    w.both_status("目标态；配置 UI 已落地列表/编辑器交互")
    w.both_h3("13.1 正常状态")
    w.both_p("展示剩余额度文案（目标态 POS/Kiosk/eMenu/SDI）。")
    w.both_h3("13.2 超限弹窗")
    w.both_p("展示规则、目标、N、L、EffectiveLimit、已下单/购物车/本次增量、超出量、可用授权范围。")
    w.both_h3("13.3 多规则批量")
    w.both_p("原子批量加购；套餐按子商品展开（目标态）。")
    w.both_h3("13.4 错误文案")
    w.both_p("须说明受限对象、原因、下一步（联系服务员授权）。")

    w.both_h1("14. 并发、幂等和版本")
    w.both_status("目标态")
    w.both_bullets([
        "orderVersion + operationId 乐观锁/事务。",
        "规则发布不可变 version；旧授权随 version 失效。",
        "客户端剩余额度仅提示，不作裁决依据。",
    ])

    # 15 split current vs target
    w.both_h1("15. 接口设计")
    w.both_h2("15.1 当前原型契约（已落地）")
    w.both_table(["类型", "键/路径", "说明"], [
        ["localStorage", "restaurantRules", "规则数组整表覆盖"],
        ["localStorage", "order-limit:rule-list-columns:v1", "列偏好"],
        ["sessionStorage", "restaurantRuleRecovery:{draftId}", "恢复副本，不读取"],
        ["路由", "/operations/queue-call/menu-order-limits", "主入口"],
        ["iframe", "order-limit.html", "列表"],
        ["iframe", "order-limit-rule-editor.html?draftId=", "六步编辑"],
        ["iframe", "order-limit-publish-confirm.html?draftId=", "发布确认"],
        ["参数", "mode=create|draftId|ruleId|copy=1|view=1|embedded=1", "见 SPEC 7.3"],
    ])
    w.both_h2("15.2 目标态 HTTP API（模板，未落地）")
    w.both_code(
        "GET/POST/PUT .../order-limit-rules\n"
        "POST .../publish|disable|archive|rollback\n"
        "POST .../limit-evaluate|limit-authorizations\n"
        "GET .../order-limit-audits\n"
        "POST .../order-limit-rules/simulate"
    )

    # 16 storage
    w.both_h1("16. 存储")
    w.both_h2("16.1 当前（已落地）")
    w.both_table(["键", "结构", "读写"], [
        ["restaurantRules", "StoredRule[]", "列表/编辑/发布"],
        ["列偏好", "{version:1,visible[]}", "列表字段设置"],
        ["恢复副本", "editorDraft JSON", "仅写/删"],
    ])
    w.both_h2("16.2 目标态表（模板，未落地）")
    w.both_bullets([
        "order_limit_rules / order_limit_rule_versions / order_limit_rule_scopes",
        "order_limit_authorizations / order_limit_audit_logs",
        "order_limit_operation_idempotency / 订单明细占用快照",
    ])

    # 17 publish
    w.both_h1("17. 发布与终端下发")
    w.both_h2("17.1 当前发布流程（已落地）")
    w.both_code(
        "步骤1-5校验 → 步骤6汇总 → 发布确认页 validateAll\n"
        "→ validateDeployStores → buildPublishedDraft\n"
        "→ status=active 写 restaurantRules → 返回列表"
    )
    w.both_h2("17.2 发布校验清单（已落地）")
    w.both_bullets([
        "人数/轮次区间连续且末段及以上。",
        "全部 limits 单元格 configured（含0）。",
        "至少1生效门店且均有商品。",
        "活动周期/会员/日期/时段完整。",
        "授权范围/默认范围/每范围权限完整。",
        "自定义确认框；禁止原生弹窗。",
    ])
    w.both_h2("17.3 终端下发（目标态）")
    w.both_p("签名规则快照推送终端；版本变化失效缓存；离线失败关闭策略见模板。")

    # 18-26
    w.both_h1("18. 旧数据迁移")
    w.both_status("已落地")
    w.both_bullets([
        "normalizeUnlimitedLimitCells：不限制→未配置，发现即写回。",
        "normalizeMergedProductQuantitySteps：七步→六步（v1/v2）。",
        "normalizeStoreDraft / normalizeDeploymentSelection。",
        "mapLegacyType/Period/Target；max>=99→null。",
        "无 storeConfigs 时从 legacy 字段生成；深拷贝隔离门店。",
    ])

    w.both_h1("19. 异常处理")
    w.both_table(["异常", "当前处理", "目标态"], [
        ["规则非法", "步骤禁用/发布阻止", "拒绝发布"],
        ["JSON 损坏", "空数组或演示数据", "—"],
        ["草稿不存在", "错误态返回", "—"],
        ["保存失败", "阻止离开", "—"],
        ["服务端不可用", "—", "终端只读"],
        ["版本冲突", "—", "刷新重试"],
        ["发布失败", "保留草稿，重新发布", "保持上一完整版本"],
    ])

    w.both_h1("20. 权限与审计")
    w.both_table(["项", "当前", "目标态"], [
        ["后台按钮 RBAC", "无，全员可操作", "查看/编辑/发布/回滚分级"],
        ["规则内所需权限", "已配置静态角色", "运行时校验员工角色"],
        ["授权审计", "无", "原限额/超限/员工/终端/范围/原因/版本"],
    ])

    w.both_h1("21. 验收标准")
    w.both_h2("21.1 配置原型（已确认，MOL-01～MOL-31）")
    w.both_bullets([
        "六步无独立商品步骤；列表滚动区固定表头；字段设置持久化。",
        "12场景可配；数量三态；活动周期/多时段/自定义时间。",
        "添加商品差异提交；完整规则行；批量/复制/预览。",
        "草稿900ms保存；session 恢复副本；发布二次确认。",
        "正式编辑不覆盖源规则；复制独立；view=1 只读。",
        "旧数据幂等迁移；自定义对话框。",
    ])
    w.both_h2("21.2 运行时（目标态，对齐模板验收清单）")
    w.both_bullets([
        "分类统计分类内总数；一分类满不影响其他分类；菜品满只限该菜品。",
        "每轮换轮重置；多轮按当前轮次区间 L；与轮次无关累计全单。",
        "按人数 EffectiveLimit=L×N；儿童策略影响 N；人数减少不追溯删历史但禁止继续加。",
        "留空≠0；分类+菜品、每轮+整单可多规则 AND 同时生效。",
        "客户端预校验与服务端最终判定一致；多终端并发不突破限额。",
        "未送厨撤销释放额度；已送厨退菜不释放。",
        "授权仅覆盖指定规则/目标/订单/范围；operation 不重复消费；round/order 换轮/关单失效。",
        "新规则 version 发布后旧授权失效；发布失败终端沿用上一完整 version。",
        "伪造人数/轮次/会员/时间/已用量不影响服务端；加购与提交互斥计数。",
        "相同 operationId+requestDigest 重试幂等；不同 digest 拒绝。",
        "批量命中多权限规则时取违反项共同允许的最小授权范围。",
        "每种 12 场景须测：小于上限、达上限、超出、0、留空及三档授权范围。",
    ])

    w.both_h1("22. 非功能指标")
    w.both_table(["指标", "配置原型", "目标态"], [
        ["判定 P95", "N/A（本地）", "<50ms 引擎 / <150ms API"],
        ["授权 P95", "N/A", "<500ms"],
        ["幂等", "N/A", "operationId+摘要"],
        ["可访问性", "tablist/对话框/a11y", "同左+终端"],
        ["验证", "verify-order-limit-*.mjs", "CI+集成测试"],
    ])

    w.both_h1("23. 成功指标")
    w.both_p("配置阶段：PRD/SPEC 与代码一致、verify 脚本通过。"
              "运行阶段（目标）：判定不一致率0、并发突破0、无审计授权0、发布失败<0.5%。")

    w.both_h1("24. 风险与缓解")
    w.both_table(["风险", "缓解"], [
        ["仅 localStorage，非生产数据源", "明确原型边界；后续 API 迁移"],
        ["旧「不限制」数据", "normalize 即时写回"],
        ["多规则难理解", "步骤6/发布预览汇总"],
        ["运行时未接入", "配置与目标态 DTO 对齐，减少二次改造"],
        ["iframe 对话框遮罩", "挂 top.document + pagehide 清理"],
    ])

    w.both_h1("25. 实施阶段与资源")
    w.both_h2("25.1 已完成（当前仓库）")
    w.both_bullets([
        "六步编辑器、列表、发布、迁移、专项 specs、PRD/SPEC、verify 脚本。",
    ])
    w.both_h2("25.2 后续（目标态，模板估算）")
    w.both_bullets([
        "规则 DTO 与服务端引擎 2-3 周；终端接入 3-4 周；联调灰度 3-4 周（参考模板）。",
    ])

    w.both_h1("26. 完成定义")
    w.both_h2("26.1 当前版本（已达成）")
    w.both_bullets([
        "12 场景配置语义与数量矩阵展开完整可用。",
        "StoredRule/EditorDraft 字段与 SPEC 一致；MOL-01～MOL-31 可达。",
        "六步发布闭环；localStorage 为唯一数据源；文档归类 menu-order-limit/。",
    ])
    w.both_h2("26.2 目标态完成定义（模板）")
    w.both_bullets([
        "12 场景运行时判定一致；服务端事务+终端下发+授权审计闭环；"
        "38 条运行时验收通过；localStorage 不再为正式数据源。",
    ])


def main() -> None:
    w = Writer()
    build(w)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text("\n".join(w.md), encoding="utf-8")
    w.doc.save(str(OUT_DOCX))
    try:
        import shutil
        shutil.copy2(OUT_DOCX, DL_DOCX)
    except OSError:
        pass
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
